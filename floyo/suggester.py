"""Integration suggestion engine."""

from typing import List, Dict, Any
from pathlib import Path
from datetime import datetime, timedelta


class IntegrationSuggester:
    """Suggests API integrations based on usage patterns."""
    
    def __init__(self, tracker):
        """Initialize the suggester.
        
        Args:
            tracker: UsageTracker instance
        """
        self.tracker = tracker
        
        # Known tool integrations
        self.tool_integrations = {
            ".py": {
                "tools": ["python", "scraper", "automation"],
                "integrations": [
                    "Dropbox API - auto-sync output files",
                    "Email API - send results via email",
                    "Database API - store results in database",
                    "Slack API - notify on completion"
                ]
            },
            ".vba": {
                "tools": ["microsoft_word", "macro"],
                "integrations": [
                    "Python script - process macro output",
                    "Dropbox API - auto-save documents",
                    "Email API - send documents automatically",
                    "File system watcher - trigger on document changes"
                ]
            },
            ".sh": {
                "tools": ["bash", "shell_script"],
                "integrations": [
                    "Cron scheduler - run automatically",
                    "Log aggregator - collect script outputs",
                    "Notification API - alert on errors",
                    "Cloud storage - backup script outputs"
                ]
            },
            ".js": {
                "tools": ["nodejs", "javascript"],
                "integrations": [
                    "Webhook API - trigger external services",
                    "Database API - store execution results",
                    "S3 API - upload generated files",
                    "Queue system - process jobs asynchronously"
                ]
            }
        }
    
    def suggest_integrations(self) -> List[Dict[str, Any]]:
        """Generate integration suggestions based on usage patterns.
        
        Returns:
            List of suggested integrations with sample code
        """
        patterns = self.tracker.get_patterns()
        recent_events = self.tracker.get_recent_events(limit=50)
        
        suggestions = []
        
        # Analyze file type patterns
        for file_ext, pattern_data in patterns.items():
            if file_ext in self.tool_integrations:
                integration_info = self.tool_integrations[file_ext]
                
                # Check if recently used
                last_used = pattern_data.get("last_used")
                if last_used:
                    last_used_dt = datetime.fromisoformat(last_used)
                    if datetime.now() - last_used_dt < timedelta(days=7):
                        suggestion = {
                            "trigger": f"Recently used {file_ext} files",
                            "tools_involved": list(pattern_data.get("tools", [])) + integration_info["tools"],
                            "suggested_integration": integration_info["integrations"][0],
                            "sample_code": self._generate_sample_code(file_ext, integration_info["integrations"][0]),
                            "reasoning": self._generate_reasoning(file_ext, pattern_data, recent_events)
                        }
                        suggestions.append(suggestion)
        
        # Look for workflow patterns across recent events
        workflow_suggestions = self._analyze_workflow_patterns(recent_events)
        suggestions.extend(workflow_suggestions)
        
        return suggestions[:5]  # Return top 5 suggestions
    
    def _analyze_workflow_patterns(self, events: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Analyze events to find workflow patterns.
        
        Args:
            events: Recent usage events
            
        Returns:
            List of workflow-based suggestions
        """
        suggestions = []
        
        # Look for patterns: file creation -> file modification -> script execution
        python_files = []
        word_files = []
        scripts = []
        
        for event in events[-20:]:  # Last 20 events
            details = event.get("details", {})
            file_path = details.get("file_path", "")
            
            if file_path.endswith(".py"):
                python_files.append((event["timestamp"], file_path))
            elif file_path.endswith(".docx") or file_path.endswith(".doc"):
                word_files.append((event["timestamp"], file_path))
            elif file_path.endswith(".sh") or file_path.endswith(".vba"):
                scripts.append((event["timestamp"], file_path))
        
        # Suggest integration if Python + Word files used recently
        if python_files and word_files:
            suggestion = {
                "trigger": "Detected Python scripts and Word documents in workflow",
                "tools_involved": ["python", "microsoft_word"],
                "suggested_integration": "Chain Python scraper output with Word document automation",
                "sample_code": self._generate_workflow_code("python_word"),
                "reasoning": "You've been working with both Python scripts and Word documents. Consider automating the flow."
            }
            suggestions.append(suggestion)
        
        return suggestions
    
    def _generate_sample_code(self, file_ext: str, integration: str) -> str:
        """Generate sample code for an integration.
        
        Args:
            file_ext: File extension
            integration: Integration description
            
        Returns:
            Sample code as string
        """
        if file_ext == ".py" and "Dropbox" in integration:
            return """# Auto-sync Python script output to Dropbox
import dropbox
from pathlib import Path

def sync_to_dropbox(local_file, remote_path):
    dbx = dropbox.Dropbox('YOUR_ACCESS_TOKEN')
    with open(local_file, 'rb') as f:
        dbx.files_upload(f.read(), remote_path)

# Use after your script generates output
output_file = 'results.json'
sync_to_dropbox(output_file, '/scripts/output/' + output_file)
"""
        
        elif file_ext == ".py" and "Email" in integration:
            return """# Send script results via email
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

def send_results_email(results_file):
    msg = MIMEMultipart()
    msg['From'] = 'your_email@example.com'
    msg['To'] = 'recipient@example.com'
    msg['Subject'] = 'Script Results'
    
    with open(results_file, 'r') as f:
        body = f.read()
    msg.attach(MIMEText(body, 'plain'))
    
    server = smtplib.SMTP('smtp.example.com', 587)
    server.starttls()
    server.login('your_email@example.com', 'password')
    server.send_message(msg)
    server.quit()

# Use after your script completes
send_results_email('output.txt')
"""
        
        elif file_ext == ".vba" and "Python" in integration:
            return r"""' VBA: Export data for Python processing
Sub ExportForPython()
    Dim outputFile As String
    outputFile = Environ("TEMP") & "\word_data.csv"
    
    ' Export current document data
    ActiveDocument.SaveAs2 FileName:=outputFile, FileFormat:=wdFormatText
    
    ' Trigger Python script (optional)
    Shell "python process_word_data.py " & outputFile, vbHide
End Sub

# Python: process_word_data.py
import csv
import sys

def process_word_data(file_path):
    with open(file_path, 'r') as f:
        # Process the exported Word data
        data = f.read()
        # Your processing logic here
        return processed_data

if __name__ == '__main__':
    result = process_word_data(sys.argv[1])
    # Save or send result
"""
        
        else:
            return f"# Integration: {integration}\n# Add your code here based on {file_ext} files"
    
    def _generate_workflow_code(self, workflow_type: str) -> str:
        """Generate code for a specific workflow.
        
        Args:
            workflow_type: Type of workflow
            
        Returns:
            Sample workflow code
        """
        if workflow_type == "python_word":
            return """# Complete workflow: Python scraper -> Word document automation
import subprocess
from docx import Document

# Step 1: Run your Python scraper
scraper_result = subprocess.run(['python', 'scraper.py'], 
                                capture_output=True, text=True)

# Step 2: Create Word document with results
doc = Document()
doc.add_heading('Scraper Results', 0)
doc.add_paragraph(scraper_result.stdout)

# Step 3: Save and optionally move to Dropbox
doc.save('scraper_results.docx')

# Step 4: Move to Dropbox (if configured)
# sync_to_dropbox('scraper_results.docx', '/reports/scraper_results.docx')
"""
        
        return f"# Workflow code for {workflow_type}"
    
    def _generate_reasoning(self, file_ext: str, pattern_data: Dict[str, Any], 
                           recent_events: List[Dict[str, Any]]) -> str:
        """Generate human-readable reasoning for a suggestion.
        
        Args:
            file_ext: File extension
            pattern_data: Pattern data for this file type
            recent_events: Recent usage events
            
        Returns:
            Reasoning text
        """
        count = pattern_data.get("count", 0)
        last_used = pattern_data.get("last_used")
        
        if last_used:
            last_used_dt = datetime.fromisoformat(last_used)
            days_ago = (datetime.now() - last_used_dt).days
            return (f"You've used {file_ext} files {count} times, "
                   f"last used {days_ago} day(s) ago. "
                   f"This integration could automate your workflow.")
        else:
            return f"You've used {file_ext} files {count} times. Consider this integration."
